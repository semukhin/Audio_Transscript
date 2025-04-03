import os
import json
import tempfile
import uuid
import time
import datetime
import threading
import re
import ssl
import urllib3
from flask import Flask, render_template, request, jsonify, send_file, url_for
from werkzeug.utils import secure_filename
import yt_dlp
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback
from config import config as app_config
import magic
from langdetect import detect, LangDetectException
from pydub import AudioSegment
from pydub.silence import detect_silence
import requests

# Отключаем проверку SSL сертификатов
ssl._create_default_https_context = ssl._create_unverified_context
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Определение конфигурации в зависимости от окружения
config_name = os.environ.get('FLASK_CONFIG', 'default')
config = app_config[config_name]

# Инициализация Flask приложения
app = Flask(__name__)
app.config.from_object(config)
app.config['SECRET_KEY'] = config.SECRET_KEY
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = config.MAX_CONTENT_LENGTH
app.config['SESSION_EXPIRY'] = config.SESSION_EXPIRY
app.config['WHISPER_SERVICE_URL'] = config.WHISPER_SERVICE_URL

# Создание папки для загрузок, если её нет
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.join(tempfile.gettempdir(), 'transcripts'), exist_ok=True)

# Разрешенные расширения файлов
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'ogg', 'flac', 'm4a', 'aac', 'opus', 'webm'}

# Словарь для хранения статусов задач и сессий
task_status = {}
sessions = {}


def generate_task_id():
    """Генерация уникального ID задачи"""
    return str(uuid.uuid4())


def generate_session_id():
    """Генерация уникального ID сессии"""
    return str(uuid.uuid4())


def allowed_file(filename):
    """Проверка допустимости расширения файла"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def format_time(seconds):
    """Форматирование времени в формат ММ:СС"""
    minutes = int(seconds) // 60
    seconds = int(seconds) % 60
    return f"{minutes:02d}:{seconds:02d}"


def check_and_convert_audio_channels(file_path, status_callback=None):
    """
    Проверяет количество каналов в аудиофайле и при необходимости конвертирует в моно
    Использует ffmpeg напрямую для более надежной обработки
    """
    try:
        import subprocess
        import json
        
        if status_callback:
            status_callback(8, "Проверка аудиофайла...")
        
        # Используем ffprobe для получения информации о файле
        cmd = [
            'ffprobe', '-v', 'quiet', '-print_format', 'json',
            '-show_streams', file_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        info = json.loads(result.stdout)
        
        # Определяем, есть ли аудио поток и сколько каналов
        audio_stream = None
        for stream in info.get('streams', []):
            if stream.get('codec_type') == 'audio':
                audio_stream = stream
                break
        
        if not audio_stream:
            if status_callback:
                status_callback(10, "Предупреждение: аудио поток не найден в файле")
            return file_path
        
        channels = audio_stream.get('channels', 0)
        
        # Если уже моно, просто возвращаем путь к файлу
        if channels == 1:
            if status_callback:
                status_callback(10, "Файл уже в монофоническом формате")
            return file_path
        
        # Конвертируем в моно, если больше одного канала
        if channels > 1:
            if status_callback:
                status_callback(9, f"Файл имеет {channels} каналов. Конвертация в моно...")
            
            output_path = f"{file_path}.mono.wav"
            cmd = [
                'ffmpeg', '-y', '-i', file_path, 
                '-ac', '1', '-ar', '16000', 
                '-vn', '-acodec', 'pcm_s16le', output_path
            ]
            
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            if status_callback:
                status_callback(10, "Файл успешно преобразован в монофонический формат")
            
            return output_path
        
        # По умолчанию возвращаем исходный файл
        return file_path
    
    except Exception as e:
        print(f"Ошибка при проверке/конвертации аудио: {e}")
        if status_callback:
            status_callback(10, f"Предупреждение: {str(e)}. Используем исходный файл.")
        return file_path


def prepare_audio_for_transcription(file_path, status_callback=None):
    """Подготовка аудиофайла для транскрипции: конвертация в подходящий формат"""
    try:
        if status_callback:
            status_callback(5, "Анализ аудиофайла...")
        
        # Получаем расширение файла
        file_ext = os.path.splitext(file_path)[1].lower()
        output_path = f"{file_path}.mono.wav"
        
        try:
            if status_callback:
                status_callback(8, "Преобразование аудио в оптимальный формат...")
            
            # Загружаем аудио с помощью pydub
            try:
                from pydub import AudioSegment
                
                if file_ext == '.mp3':
                    audio = AudioSegment.from_mp3(file_path)
                elif file_ext == '.m4a':
                    audio = AudioSegment.from_file(file_path, format="m4a")
                elif file_ext == '.ogg':
                    audio = AudioSegment.from_ogg(file_path)
                elif file_ext == '.flac':
                    audio = AudioSegment.from_file(file_path, format="flac")
                elif file_ext == '.webm':
                    audio = AudioSegment.from_file(file_path, format="webm")
                elif file_ext == '.opus':
                    audio = AudioSegment.from_file(file_path, format="opus")
                elif file_ext == '.wav':
                    audio = AudioSegment.from_wav(file_path)
                else:
                    audio = AudioSegment.from_file(file_path)
                
                # Всегда преобразуем в моно, независимо от исходного формата
                audio = audio.set_channels(1)  # 1 канал (моно)
                audio = audio.set_sample_width(2)  # 16 bit
                audio = audio.set_frame_rate(16000)  # 16 kHz частота дискретизации
                
                # Сохраняем в моно WAV
                audio.export(output_path, format="wav")
                
                if status_callback:
                    status_callback(10, "Аудио успешно преобразовано в формат WAV mono для распознавания")
                
                return output_path
            except ImportError:
                # Если pydub не установлен, используем ffmpeg
                raise ImportError("pydub не установлен, используем ffmpeg")
                
        except Exception as e:
            print(f"Ошибка при конвертации аудио: {e}")
            traceback.print_exc()
            if status_callback:
                status_callback(10, f"Ошибка конвертации: {str(e)}. Попытка конвертации с другими параметрами.")
            
            # Попытка аварийной конвертации через ffmpeg
            try:
                import subprocess
                cmd = [
                    'ffmpeg', '-y', '-i', file_path, 
                    '-ac', '1', '-ar', '16000', 
                    '-vn', '-acodec', 'pcm_s16le', output_path
                ]
                subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                
                if status_callback:
                    status_callback(10, "Аудио успешно преобразовано через ffmpeg")
                
                return output_path
            except Exception as e2:
                print(f"Ошибка при использовании ffmpeg: {e2}")
                if status_callback:
                    status_callback(10, f"Ошибка резервной конвертации: {str(e2)}. Используем исходный файл.")
                return file_path
    
    except ImportError:
        # Если pydub не установлен
        print("Предупреждение: библиотека pydub не установлена. Используем ffmpeg напрямую.")
        try:
            # Попытка использовать ffmpeg напрямую
            import subprocess
            output_path = f"{file_path}.mono.wav"
            cmd = [
                'ffmpeg', '-y', '-i', file_path, 
                '-ac', '1', '-ar', '16000', 
                '-vn', '-acodec', 'pcm_s16le', output_path
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            if status_callback:
                status_callback(10, "Аудио успешно преобразовано через ffmpeg")
            
            return output_path
        except Exception as e:
            print(f"Ошибка при использовании ffmpeg: {e}")
            return file_path


def check_audio_for_speech(file_path, status_callback=None):
    """Проверка аудиофайла на наличие речи и шумов"""
    try:
        if status_callback:
            status_callback(12, "Проверка аудио на наличие речи...")
        
        try:
            from pydub import AudioSegment
            from pydub.silence import detect_nonsilent
            
            audio = AudioSegment.from_file(file_path)
            
            # Проверяем на наличие не-тишины
            non_silent_parts = detect_nonsilent(audio, min_silence_len=500, silence_thresh=-40)
            
            if not non_silent_parts:
                if status_callback:
                    status_callback(15, "В аудиофайле не обнаружена речь (только тишина)")
                return False, "В аудиофайле не обнаружена речь (только тишина)"
            
            # Проверяем среднюю громкость
            if audio.dBFS < -30:
                if status_callback:
                    status_callback(15, "Аудиофайл имеет очень низкую громкость")
                return True, "Аудиофайл имеет очень низкую громкость, распознавание может быть неточным"
            
            if status_callback:
                status_callback(15, "Аудиофайл содержит речь с нормальным уровнем громкости")
            return True, "Аудиофайл содержит речь"
        except Exception as e:
            print(f"Ошибка при проверке аудио: {e}")
            return True, "Не удалось проверить аудиофайл на наличие речи"
    except ImportError:
        # Если pydub не установлен
        return True, "Предупреждение: библиотека pydub не установлена. Проверка аудио недоступна."


def detect_speaker_names(transcript):
    """
    Попытка определить имена говорящих из контекста разговора
    Это упрощенная версия - в реальной системе нужен более сложный алгоритм NLP
    """
    if isinstance(transcript, list):
        # Ищем имена в формате "Имя:"
        name_pattern = re.compile(r'([А-Я][а-я]+):', re.UNICODE)
        potential_names = set()
        
        for segment in transcript:
            if 'text' in segment:
                matches = name_pattern.findall(segment['text'])
                for match in matches:
                    if len(match) > 2:  # Фильтруем слишком короткие "имена"
                        potential_names.add(match)
        
        # Если нашли имена, заменяем "Говорящий X" на имена
        if potential_names:
            names_list = list(potential_names)
            speaker_map = {}
            
            # Создаем соответствие "Говорящий X" -> "Имя"
            for i, segment in enumerate(transcript):
                if 'speaker' in segment and segment['speaker'].startswith('Говорящий '):
                    speaker_num = segment['speaker'].split(' ')[1]
                    if speaker_num.isdigit():
                        idx = int(speaker_num) - 1
                        if idx < len(names_list) and speaker_num not in speaker_map:
                            speaker_map[speaker_num] = names_list[idx]
            
            # Заменяем имена говорящих
            for segment in transcript:
                if 'speaker' in segment and segment['speaker'].startswith('Говорящий '):
                    speaker_num = segment['speaker'].split(' ')[1]
                    if speaker_num in speaker_map:
                        segment['speaker'] = speaker_map[speaker_num]
    
    return transcript


def detect_audio_language(file_path):
    """Определение языка аудио с улучшенной логикой"""
    try:
        # Сначала проверяем, указан ли язык в метаданных или имени файла
        basename = os.path.basename(file_path).lower()
        
        # Проверка на явные указания языка в имени файла
        if any(lang in basename for lang in ['рус', 'rus', 'russian', '.ru']):
            return 'ru-RU'
        if any(lang in basename for lang in ['eng', 'english', '.en']):
            return 'en-US'
        
        # По умолчанию для большинства случаев используем русский язык
        return 'ru-RU'  # Всегда возвращаем русский язык по умолчанию
    except Exception as e:
        print(f"Ошибка при определении языка аудио: {e}")
        return 'ru-RU'  # В случае ошибки также используем русский


def split_audio_on_silence(file_path, min_silence_len=700, silence_thresh=-40, 
                         min_segment_len=45000, max_segment_len=55000,
                         pause_search_start=50000, pause_search_end=58000):
    """
    Разделяет аудиофайл на сегменты с интеллектуальным поиском пауз.

    Args:
        file_path: путь к аудиофайлу
        min_silence_len: минимальная длина тишины (мс)
        silence_thresh: порог тишины (дБ)
        min_segment_len: минимальная длина сегмента (мс)
        max_segment_len: максимальная длина сегмента (мс)
        pause_search_start: начало диапазона поиска паузы (мс)
        pause_search_end: конец диапазона поиска паузы (мс)
    
    Returns:
        Список кортежей (начало, конец) для каждого сегмента в миллисекундах
    """
    try:
        audio = AudioSegment.from_file(file_path)
        audio_len = len(audio)
        segments = []
        start = 0

        while start < audio_len:
            # Определяем конец текущего сегмента
            end = min(start + max_segment_len, audio_len)
            
            # Если это последний фрагмент или он короткий, сохраняем как есть
            if end - start <= min_segment_len or end == audio_len:
                segments.append((start, end))
                break
            
            # Ищем подходящую паузу для разделения
            search_end = min(start + pause_search_end, end)
            search_start = max(start + pause_search_start, start + min_segment_len)
            
            segment = audio[search_start:search_end]
            silence_ranges = detect_silence(segment, 
                                        min_silence_len=min_silence_len,
                                        silence_thresh=silence_thresh)
            
            if silence_ranges:
                # Берем середину самой длинной паузы
                longest_silence = max(silence_ranges, key=lambda x: x[1] - x[0])
                split_point = search_start + (longest_silence[0] + longest_silence[1]) // 2
                segments.append((start, split_point))
                start = split_point
            else:
                # Если паузу не нашли, делим по максимальной длине
                segments.append((start, end))
                start = end

        return segments
    except Exception as e:
        print(f"Ошибка при разделении аудио: {e}")
        # Возвращаем весь файл как один сегмент в случае ошибки
        return [(0, audio_len)] if 'audio_len' in locals() else []


# Импорт whisper_client для взаимодействия с новым сервисом
from whisper_client import transcribe_with_whisper_api

def transcribe_audio(file_path, language_code='ru-RU', enable_timestamps=False, status_callback=None):
    """Переработанная функция транскрибирования с использованием нового Whisper API"""
    def update_status(percent, message):
        print(f"[Прогресс] {percent}%: {message}")
        if status_callback:
            status_callback(percent, message)

    try:
        # Подготовка аудио
        update_status(10, "Подготовка аудиофайла...")
        file_path = check_and_convert_audio_channels(file_path, update_status)
        prepared_file_path = prepare_audio_for_transcription(file_path, update_status)

        # Проверка на наличие речи
        has_speech, speech_message = check_audio_for_speech(prepared_file_path, update_status)
        if not has_speech:
            return speech_message

        # Транскрипция через новый Whisper API
        update_status(30, "Отправка файла на транскрипцию (Whisper Russian)...")
        
        transcript = transcribe_with_whisper_api(
            prepared_file_path,
            language_code=language_code,
            enable_timestamps=enable_timestamps,
            status_callback=update_status
        )
        
        # Очистка временных файлов
        if prepared_file_path != file_path and os.path.exists(prepared_file_path):
            try:
                os.remove(prepared_file_path)
            except Exception as e:
                print(f"Не удалось удалить временный файл: {e}")
        
        # Если получили массив с таймкодами, обрабатываем имена говорящих
        if enable_timestamps and isinstance(transcript, list):
            transcript = detect_speaker_names(transcript)

        return transcript

    except Exception as e:
        print(f"Ошибка при транскрибировании: {e}")
        traceback.print_exc()
        return f"Ошибка при транскрибировании: {str(e)}"


def get_video_info(url):
    """Получение информации о видео по ссылке"""
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': '%(title)s.%(ext)s',  # Временный шаблон имени файла
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'wav',
            'preferredquality': '192',
        }],
        'quiet': False,
        'no_warnings': False,
        'geo_bypass': True,
        'nocheckcertificate': True,
        'ssl_verify': False,
        'retries': 3,
        'extractor_args': {
            'youtube': {
                'skip': ['dash', 'hls'],
                'player_skip': ['js', 'configs', 'webpage']
            }
        },
        'progress_hooks': [],
        'verify': False,  # Отключаем проверку SSL
        'no_check_certificate': True,  # Дополнительное отключение проверки сертификатов
        'legacyserverconnect': True,  # Используем устаревший метод подключения
        'no_warnings': True,
        'quiet': True,
        'socket_timeout': 30,  # Увеличиваем таймаут
        'source_address': '0.0.0.0'  # Привязываем к любому IP
    }
    
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            return {
                'title': info.get('title', 'Неизвестное видео'),
                'uploader': info.get('uploader', 'Неизвестный автор'),
                'duration': info.get('duration', 0),
                'upload_date': info.get('upload_date', ''),
                'thumbnail': info.get('thumbnail', '')
            }
    except Exception as e:
        print(f"Ошибка при получении информации о видео: {e}")
        return None


def download_from_youtube(url, status_callback=None):
    """Загрузка аудио из YouTube видео"""
    try:
        if status_callback:
            status_callback(5, "Подготовка к загрузке видео...")
        
        # Создаем временную директорию для загрузки
        temp_dir = tempfile.mkdtemp()
        temp_file = os.path.join(temp_dir, 'audio')
        
        # Настройки для yt-dlp с полностью отключенной проверкой SSL
        ydl_opts = {
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'wav',
                'preferredquality': '192',
            }],
            'outtmpl': temp_file,
            'nocheckcertificate': True,
            'no_warnings': True,
            'quiet': True,
            'extract_flat': False,
            'force_generic_extractor': False,
            'ssl_verify': False,
            'geo_bypass': True,
            'retries': 3,
            'extractor_args': {
                'youtube': {
                    'skip': ['dash', 'hls'],
                    'player_skip': ['js', 'configs', 'webpage']
                }
            },
            'progress_hooks': [lambda d: download_progress_hook(d, status_callback) if status_callback else None],
            'verify': False,  # Отключаем проверку SSL
            'no_check_certificate': True,  # Дополнительное отключение проверки сертификатов
            'legacyserverconnect': True,  # Используем устаревший метод подключения
            'no_warnings': True,
            'quiet': True
        }
        
        def download_progress_hook(d, callback):
            if d['status'] == 'downloading':
                try:
                    percent = int(float(d['_percent_str'].replace('%', '').strip()))
                    callback(5 + int(percent * 0.3), f"Загрузка видео: {percent}%")
                except:
                    pass
            elif d['status'] == 'finished':
                callback(35, "Загрузка завершена, извлечение аудио...")
        
        if status_callback:
            status_callback(5, "Получение информации о видео...")
        
        # Загружаем видео
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            try:
                # Получаем информацию о видео
                info = ydl.extract_info(url, download=True)
                
                # Ищем сконвертированный WAV файл
                wav_file = temp_file + '.wav'
                if not os.path.exists(wav_file):
                    # Если WAV не найден, ищем другие аудиофайлы
                    for ext in ['.mp3', '.m4a', '.opus']:
                        if os.path.exists(temp_file + ext):
                            wav_file = temp_file + ext
                            break
                
                if not os.path.exists(wav_file):
                    raise Exception("Не удалось найти загруженный аудиофайл")
                
                if status_callback:
                    status_callback(40, "Аудио успешно извлечено")
                
                # Возвращаем путь к файлу и информацию о видео
                return wav_file, {
                    'title': info.get('title', 'Неизвестное видео'),
                    'uploader': info.get('uploader', 'Неизвестный автор'),
                    'duration': info.get('duration', 0),
                    'description': info.get('description', ''),
                    'upload_date': info.get('upload_date', ''),
                }
                
            except Exception as e:
                if status_callback:
                    status_callback(0, f"Ошибка при загрузке видео: {str(e)}")
                raise Exception(f"Ошибка при загрузке видео: {str(e)}")
    
    except Exception as e:
        if status_callback:
            status_callback(0, f"Ошибка: {str(e)}")
        raise Exception(f"Ошибка при загрузке видео: {str(e)}")


def create_docx(transcript, filename="transcript", with_timestamps=False, video_info=None):
    """Создание DOCX файла с транскрипцией."""
    doc = docx.Document()
    
    # Стилизация заголовка
    title = doc.add_heading("Транскрипция аудио", 0)
    title_paragraph = title.paragraph_format
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Добавление информации о видео, если она есть
    if video_info:
        video_paragraph = doc.add_paragraph()
        video_paragraph.add_run("Информация о видео:\n").bold = True
        video_paragraph.add_run(f"Название: {video_info['title']}\n")
        video_paragraph.add_run(f"Автор: {video_info['uploader']}\n")
        
        if video_info['duration']:
            minutes, seconds = divmod(video_info['duration'], 60)
            video_paragraph.add_run(f"Длительность: {minutes}:{seconds:02d}\n")
        
        if video_info['upload_date']:
            upload_date = video_info['upload_date']
            formatted_date = f"{upload_date[6:8]}.{upload_date[4:6]}.{upload_date[0:4]}"
            video_paragraph.add_run(f"Дата публикации: {formatted_date}\n")
    
    # Добавление даты и времени
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Дата транскрипции: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    date_run.font.size = Pt(10)
    date_run.font.italic = True
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Горизонтальная линия
    doc.add_paragraph("_" * 80)
    
    # Добавление транскрибированного текста с красивым оформлением
    if with_timestamps and isinstance(transcript, list):
        # Оформление с таймингами и разными говорящими
        for segment in transcript:
            paragraph = doc.add_paragraph()
            
            # Добавление тайминга
            time_run = paragraph.add_run(f"[{segment['start_time']}] ")
            time_run.font.bold = True
            time_run.font.size = Pt(10)
            time_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Добавление говорящего
            speaker_run = paragraph.add_run(f"{segment['speaker']}: ")
            speaker_run.font.bold = True
            speaker_run.font.color.rgb = RGBColor(0, 0, 150)
            text_run = paragraph.add_run(segment['text'])
            text_run.font.size = Pt(11)
            
            # Добавление пустой строки между говорящими
            doc.add_paragraph()
    else:
        # Обычный текст без разделения на говорящих
        paragraph = doc.add_paragraph()
        text_run = paragraph.add_run(transcript)
        text_run.font.size = Pt(11)
    
    # Добавление нижнего колонтитула
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Создано с помощью сервиса транскрипции аудио (Whisper Russian)")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    
    # Сохранение во временный файл
    temp_dir = os.path.join(tempfile.gettempdir(), 'transcripts')
    os.makedirs(temp_dir, exist_ok=True)
    temp_file = os.path.join(temp_dir, f"{filename}.docx")
    doc.save(temp_file)
    
    return temp_file


def save_transcript_to_session(session_id, transcript, docx_path, with_timestamps=False, video_info=None, language_code='ru-RU'):
    """Сохранение транскрипции в сессию с информацией о языке"""
    # Формируем уникальный URL для доступа к сессии
    share_url = f"/share/{session_id}"
    
    sessions[session_id] = {
        'created_at': datetime.datetime.now().timestamp(),
        'transcript': transcript,
        'docx_path': docx_path,
        'with_timestamps': with_timestamps,
        'video_info': video_info,
        'share_url': share_url,
        'language': language_code  # Добавляем информацию о языке
    }
    
    # Очистка старых сессий (старше 24 часов)
    current_time = datetime.datetime.now().timestamp()
    sessions_to_delete = []
    
    for s_id, s_data in sessions.items():
        if current_time - s_data['created_at'] > app.config['SESSION_EXPIRY']:
            sessions_to_delete.append(s_id)
    
    for s_id in sessions_to_delete:
        # Также удаляем соответствующие DOCX файлы
        if 'docx_path' in sessions[s_id]:
            docx_file = os.path.join(tempfile.gettempdir(), 'transcripts', sessions[s_id]['docx_path'])
            try:
                if os.path.exists(docx_file):
                    os.remove(docx_file)
            except Exception as e:
                print(f"Ошибка при удалении устаревшего DOCX файла: {e}")
        del sessions[s_id]
    
    return share_url


def analyze_transcript(transcript, with_timestamps=False):
    """Расширенный анализ транскрипции с дополнительными метриками"""
    try:
        result = {}
        
        if with_timestamps and isinstance(transcript, list):
            # Статистика по говорящим
            speakers = {}
            word_frequencies = {}
            sentence_lengths = []
            total_words = 0
            total_time = 0  # Общая длительность в секундах
            
            for segment in transcript:
                speaker = segment['speaker']
                text = segment['text']
                
                # Преобразование времени (format: "MM:SS")
                start_time_str = segment.get('start_time', '00:00')
                try:
                    minutes, seconds = map(int, start_time_str.split(':'))
                    start_time_sec = minutes * 60 + seconds
                except (ValueError, AttributeError):
                    start_time_sec = 0
                
                # Разбиваем текст на слова и предложения
                words = re.findall(r'\b[а-яА-Яa-zA-Z]+\b', text.lower())
                sentences = re.split(r'[.!?]+', text)
                
                # Подсчет для говорящего
                if speaker not in speakers:
                    speakers[speaker] = {
                        'word_count': 0,
                        'sentence_count': 0,
                        'total_time': 0,
                        'avg_words_per_sentence': 0,
                        'speech_rate': 0,  # Слов в минуту
                        'segments': 0
                    }
                
                speakers[speaker]['word_count'] += len(words)
                speakers[speaker]['sentence_count'] += len(sentences)
                speakers[speaker]['segments'] += 1
                
                # Оценка времени (очень приблизительно)
                if len(transcript) > 1:  # Если есть несколько сегментов
                    # Находим следующий сегмент для того же говорящего
                    next_segment_idx = -1
                    for i, seg in enumerate(transcript):
                        if seg == segment:  # Нашли текущий сегмент
                            # Ищем следующий с тем же говорящим
                            for j in range(i+1, len(transcript)):
                                if transcript[j]['speaker'] == speaker:
                                    next_segment_idx = j
                                    break
                            break
                    
                    if next_segment_idx > 0:
                        # Найден следующий сегмент того же говорящего
                        next_start_time_str = transcript[next_segment_idx].get('start_time', '00:00')
                        try:
                            next_minutes, next_seconds = map(int, next_start_time_str.split(':'))
                            next_start_time_sec = next_minutes * 60 + next_seconds
                            segment_duration = next_start_time_sec - start_time_sec
                            if segment_duration > 0:
                                speakers[speaker]['total_time'] += segment_duration
                                total_time += segment_duration
                        except (ValueError, AttributeError):
                            pass
                
                # Подсчет общей частоты слов
                for word in words:
                    if len(word) > 2:  # Игнорируем слишком короткие слова
                        word_frequencies[word] = word_frequencies.get(word, 0) + 1
                
                # Длина предложений
                for sentence in sentences:
                    if sentence.strip():
                        words_in_sentence = len(re.findall(r'\b[а-яА-Яa-zA-Z]+\b', sentence.lower()))
                        if words_in_sentence > 0:
                            sentence_lengths.append(words_in_sentence)
                
                total_words += len(words)
            
            # Расчет среднего времени между сегментами (если доступно)
            average_segment_duration = 0
            if len(transcript) > 1:
                # Сортируем сегменты по времени
                sorted_segments = sorted(transcript, key=lambda x: x.get('start_time', '00:00'))
                total_duration = 0
                segments_with_time = 0
                
                for i in range(1, len(sorted_segments)):
                    current_time_str = sorted_segments[i].get('start_time', '')
                    prev_time_str = sorted_segments[i-1].get('start_time', '')
                    
                    if current_time_str and prev_time_str:
                        try:
                            # Вычисляем разницу времени
                            current_minutes, current_seconds = map(int, current_time_str.split(':'))
                            current_time_sec = current_minutes * 60 + current_seconds
                            
                            prev_minutes, prev_seconds = map(int, prev_time_str.split(':'))
                            prev_time_sec = prev_minutes * 60 + prev_seconds
                            
                            duration = current_time_sec - prev_time_sec
                            if duration > 0:
                                total_duration += duration
                                segments_with_time += 1
                        except (ValueError, AttributeError):
                            pass
                
                if segments_with_time > 0:
                    average_segment_duration = total_duration / segments_with_time
            
            # Расчет средних значений
            for speaker, stats in speakers.items():
                if stats['sentence_count'] > 0:
                    stats['avg_words_per_sentence'] = round(stats['word_count'] / stats['sentence_count'], 2)
                
                if stats['total_time'] > 0:
                    # Расчет скорости речи (слов в минуту)
                    stats['speech_rate'] = round(stats['word_count'] / (stats['total_time'] / 60), 1)
            
            # Выявление ключевых фраз (простой алгоритм)
            keywords = {}
            for word, count in word_frequencies.items():
                if count > 2 and len(word) > 3:  # Минимальный порог для ключевых слов
                    keywords[word] = count
            
            # Формирование "облака тегов"
            sorted_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
            top_keywords = dict(sorted_keywords[:30])
            
            # Сортируем частоты слов
            sorted_words = sorted(word_frequencies.items(), key=lambda x: x[1], reverse=True)
            top_words = dict(sorted_words[:50])
            
            # Определение языка транскрипции
            language = "unknown"
            try:
                # Объединяем весь текст для более точного определения
                full_text = " ".join([segment['text'] for segment in transcript])
                if full_text.strip():
                    language = detect(full_text)
            except (ImportError, LangDetectException, Exception) as e:
                print(f"Не удалось определить язык: {e}")
            
            # Формируем результат
            result = {
                'speakers': speakers,
                'top_words': top_words,
                'keywords': top_keywords,
                'total_words': total_words,
                'avg_sentence_length': round(sum(sentence_lengths) / len(sentence_lengths), 2) if sentence_lengths else 0,
                'sentence_count': len(sentence_lengths),
                'word_variety': round(len(word_frequencies) / total_words, 3) if total_words > 0 else 0,
                'avg_segment_duration': round(average_segment_duration, 2) if average_segment_duration > 0 else 0,
                'estimated_total_duration': total_time,
                'language': language
            }
        else:
            # Анализ для обычного текста без таймингов
            words = re.findall(r'\b[а-яА-Яa-zA-Z]+\b', transcript.lower())
            sentences = re.split(r'[.!?]+', transcript)
            
            # Подсчет частоты слов
            word_frequencies = {}
            for word in words:
                if len(word) > 2:
                    word_frequencies[word] = word_frequencies.get(word, 0) + 1
            
            # Сортируем частоты слов
            sorted_words = sorted(word_frequencies.items(), key=lambda x: x[1], reverse=True)
            top_words = dict(sorted_words[:50])
            
            # Выявление ключевых фраз
            keywords = {}
            for word, count in word_frequencies.items():
                if count > 2 and len(word) > 3:
                    keywords[word] = count
            
            sorted_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
            top_keywords = dict(sorted_keywords[:30])
            
            # Длина предложений
            sentence_lengths = []
            for sentence in sentences:
                if sentence.strip():
                    words_in_sentence = len(re.findall(r'\b[а-яА-Яa-zA-Z]+\b', sentence.lower()))
                    if words_in_sentence > 0:
                        sentence_lengths.append(words_in_sentence)
            
            # Определение языка транскрипции
            language = "unknown"
            try:
                if transcript.strip():
                    language = detect(transcript)
            except (ImportError, LangDetectException, Exception) as e:
                print(f"Не удалось определить язык: {e}")
            
            # Формируем результат
            result = {
                'top_words': top_words,
                'keywords': top_keywords,
                'total_words': len(words),
                'avg_sentence_length': round(sum(sentence_lengths) / len(sentence_lengths), 2) if sentence_lengths else 0,
                'sentence_count': len(sentence_lengths),
                'word_variety': round(len(word_frequencies) / len(words), 3) if words else 0,
                'language': language
            }
        
        return result
    except Exception as e:
        print(f"Ошибка при анализе транскрипции: {e}")
        traceback.print_exc()
        return {
            'status': 'error',
            'message': f'Ошибка при анализе: {str(e)}'
        }


def check_transcription_quality(transcript):
    """
    Проверка качества транскрипции
    """
    if isinstance(transcript, list):
        # Проверка на повторы
        unique_texts = set(segment['text'] for segment in transcript)
        
        # Проверка общей длины и информативности
        total_words = sum(len(segment['text'].split()) for segment in transcript)
        
        # Оценка качества
        if len(unique_texts) < len(transcript) * 0.5:
            return "Возможно дублирование текста"
        
        if total_words < 10:
            return "Слишком короткая транскрипция"
        
        return "Транскрипция выглядит корректной"
    elif isinstance(transcript, str):
        words = transcript.split()
        if len(words) < 10:
            return "Слишком короткая транскрипция"
        return "Транскрипция выглядит корректной"
    
    return "Нет данных для анализа"


# Определение маршрутов веб-приложения
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/share/<session_id>')
def shared_transcript(session_id):
    """Страница с доступом к сохраненной транскрипции"""
    if session_id in sessions:
        session_data = sessions[session_id]
        return render_template(
            'shared.html',
            transcript=session_data['transcript'],
            with_timestamps=session_data['with_timestamps'],
            video_info=session_data['video_info'],
            docx_path=session_data['docx_path'],
            language=session_data.get('language', 'ru-RU')  # Передаем язык в шаблон
        )
    return render_template('error.html', message="Сессия не найдена или истекла")


@app.route('/task_status/<task_id>', methods=['GET'])
def get_task_status(task_id):
    """Получение статуса задачи по ID"""
    if task_id in task_status:
        return jsonify(task_status[task_id])
    return jsonify({'status': 'unknown', 'percent': 0, 'message': 'Задача не найдена'})


def process_audio_file(file_path, enable_timestamps, task_id, language_code='ru-RU'):
    """Обработка аудиофайла в отдельном потоке"""
    try:
        # Функция обновления статуса
        def update_status(percent, message):
            task_status[task_id] = {
                'status': 'transcribing' if percent < 100 else 'complete',
                'percent': percent,
                'message': message
            }
        
        # Обновляем начальный статус
        update_status(5, "Начало транскрибирования с улучшенной русской моделью Whisper")
        
        # Переключаемся на использование новой функции транскрибирования
        transcript = transcribe_audio(
            file_path, 
            language_code=language_code,
            enable_timestamps=enable_timestamps, 
            status_callback=update_status
        )
        
        # Генерируем ID сессии
        session_id = generate_session_id()
        
        # Создание DOCX файла
        filename = os.path.splitext(os.path.basename(file_path))[0]
        docx_path = create_docx(transcript, filename, with_timestamps=enable_timestamps)
        
        # Сохраняем транскрипцию в сессию
        share_url = save_transcript_to_session(
            session_id, 
            transcript, 
            os.path.basename(docx_path), 
            enable_timestamps,
            language_code=language_code
        )
        
        # Финальное обновление статуса
        task_status[task_id] = {
            'status': 'complete',
            'percent': 100,
            'message': 'Транскрипция завершена',
            'transcript': transcript,
            'docx_path': os.path.basename(docx_path),
            'with_timestamps': enable_timestamps,
            'session_id': session_id,
            'share_url': share_url,
            'language': language_code
        }
        
        # Удаляем исходный аудиофайл, если транскрипция успешно завершена
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Ошибка при удалении исходного аудиофайла: {e}")
            
    except Exception as e:
        print(f"Ошибка при обработке файла: {e}")
        traceback.print_exc()
        # В случае ошибки
        task_status[task_id] = {
            'status': 'error',
            'percent': 0,
            'message': f'Ошибка: {str(e)}'
        }
        
        
def process_youtube_link(url, enable_timestamps, task_id, language_code='ru-RU'):
    """Обработка ссылки на YouTube в отдельном потоке с использованием Whisper"""
    try:
        # Функция обновления статуса
        def update_status(percent, message):
            task_status[task_id] = {
                'status': 'transcribing' if percent < 100 else 'complete',
                'percent': percent,
                'message': message
            }
        
        # Обновляем начальный статус
        update_status(5, "Начало обработки ссылки")
        
        # Загрузка аудио из видео
        audio_path, video_info = download_from_youtube(url, update_status)
        
        if not audio_path:
            task_status[task_id] = {
                'status': 'error',
                'percent': 0,
                'message': 'Не удалось загрузить аудио по указанной ссылке'
            }
            return
        
        # Запускаем транскрибирование с указанным языком
        transcript = transcribe_audio(
            audio_path, 
            enable_timestamps=enable_timestamps, 
            status_callback=update_status,
            language_code=language_code
        )
        
        # Генерируем ID сессии
        session_id = generate_session_id()
        
        # Создание DOCX файла
        docx_path = create_docx(transcript, f"link_{uuid.uuid4()}", with_timestamps=enable_timestamps, video_info=video_info)
        
        # Сохраняем транскрипцию в сессию
        share_url = save_transcript_to_session(
            session_id, 
            transcript, 
            os.path.basename(docx_path), 
            enable_timestamps, 
            video_info,
            language_code=language_code
        )
        
        # Удаление временного файла
        try:
            if os.path.exists(audio_path):
                os.remove(audio_path)
        except Exception as e:
            print(f"Ошибка при удалении временного файла: {e}")
        
        # Финальное обновление статуса
        task_status[task_id] = {
            'status': 'complete',
            'percent': 100,
            'message': 'Транскрипция завершена',
            'transcript': transcript,
            'docx_path': os.path.basename(docx_path),
            'with_timestamps': enable_timestamps,
            'video_info': video_info,
            'session_id': session_id,
            'share_url': share_url,
            'language': language_code
        }
    except Exception as e:
        print(f"Ошибка при обработке ссылки: {e}")
        traceback.print_exc()
        # В случае ошибки
        task_status[task_id] = {
            'status': 'error',
            'percent': 0,
            'message': f'Ошибка: {str(e)}'
        }


@app.route('/upload', methods=['POST'])
def upload_file():
    """Обработка загрузки аудиофайла."""
    if 'file' not in request.files:
        return jsonify({'error': 'Файл не найден в запросе'}), 400
    
    file = request.files['file']
    enable_timestamps = request.form.get('timestamps') == 'true'
    # Получаем выбранный язык или используем русский по умолчанию
    language_code = request.form.get('language', 'ru-RU')
    
    if file.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{filename}")
        file.save(file_path)
        
        # Создаем ID задачи
        task_id = generate_task_id()
        
        # Инициализируем статус задачи
        task_status[task_id] = {
            'status': 'preparing',
            'percent': 0,
            'message': 'Подготовка к обработке файла'
        }
        
        # Запускаем обработку в отдельном потоке, передавая язык
        threading.Thread(
            target=process_audio_file, 
            args=(file_path, enable_timestamps, task_id, language_code),
            daemon=True
        ).start()
        
        # Возвращаем ID задачи клиенту
        return jsonify({
            'task_id': task_id
        })
    
    return jsonify({'error': 'Формат файла не поддерживается'}), 400


@app.route('/record', methods=['POST'])
def process_recording():
    """Обработка записанного аудио."""
    if 'audio_data' not in request.files:
        return jsonify({'error': 'Аудиоданные не найдены в запросе'}), 400
    
    audio_file = request.files['audio_data']
    enable_timestamps = request.form.get('timestamps') == 'true'
    # Получаем выбранный язык или используем русский по умолчанию
    language_code = request.form.get('language', 'ru-RU')
    
    if audio_file.filename == '':
        return jsonify({'error': 'Файл не выбран'}), 400
    
    filename = f"recording_{uuid.uuid4()}.wav"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    audio_file.save(file_path)
    
    # Создаем ID задачи
    task_id = generate_task_id()
    
    # Инициализируем статус задачи
    task_status[task_id] = {
        'status': 'preparing',
        'percent': 0,
        'message': 'Подготовка к обработке записи'
    }
    
    # Запускаем обработку в отдельном потоке, передавая язык
    threading.Thread(
        target=process_audio_file, 
        args=(file_path, enable_timestamps, task_id, language_code),
        daemon=True
    ).start()
    
    # Возвращаем ID задачи клиенту
    return jsonify({
        'task_id': task_id
    })


@app.route('/link', methods=['POST'])
def process_link():
    """Обработка ссылки на видео."""
    data = request.json
    
    if not data or 'url' not in data:
        return jsonify({'error': 'URL не найден в запросе'}), 400
    
    url = data['url']
    enable_timestamps = data.get('timestamps', False)
    # Получаем выбранный язык или используем русский по умолчанию
    language_code = data.get('language', 'ru-RU')
    
    # Создаем ID задачи
    task_id = generate_task_id()
    
    # Инициализируем статус задачи
    task_status[task_id] = {
        'status': 'preparing',
        'percent': 0,
        'message': 'Подготовка к загрузке видео'
    }
    
    # Запускаем обработку в отдельном потоке, передавая язык
    threading.Thread(
        target=process_youtube_link, 
        args=(url, enable_timestamps, task_id, language_code),
        daemon=True
    ).start()
    
    # Возвращаем ID задачи клиенту
    return jsonify({
        'task_id': task_id
    })


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Скачивание DOCX файла с транскрипцией."""
    file_path = os.path.join(tempfile.gettempdir(), 'transcripts', filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'Файл не найден'}), 404
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=filename
    )


@app.route('/api/verify_link', methods=['POST'])
def verify_link():
    """API для проверки ссылки на возможность загрузки"""
    data = request.json
    
    if not data or 'url' not in data:
        return jsonify({'error': 'URL не найден в запросе'}), 400
    
    url = data['url']
    
    try:
        video_info = get_video_info(url)
        if video_info:
            return jsonify({
                'status': 'success',
                'video_info': video_info
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Не удалось получить информацию о видео'
            })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка при проверке ссылки: {str(e)}'
        })


@app.route('/api/analyze', methods=['POST'])
def analyze_transcript_api():
    """API для анализа транскрипции"""
    data = request.json
    
    if not data or 'transcript' not in data:
        return jsonify({'error': 'Транскрипция не найдена в запросе'}), 400
    
    transcript = data['transcript']
    with_timestamps = data.get('with_timestamps', False)
    
    try:
        # Используем улучшенную функцию анализа
        result = analyze_transcript(transcript, with_timestamps)
        
        return jsonify({
            'status': 'success',
            'analysis': result
        })
    except Exception as e:
        print(f"Ошибка при анализе транскрипции: {e}")
        traceback.print_exc()
        return jsonify({
            'status': 'error',
            'message': f'Ошибка при анализе: {str(e)}'
        }), 500


@app.route('/health', methods=['GET'])
def health_check():
    """Проверка работоспособности сервиса"""
    return jsonify({
        'status': 'healthy',
        'version': '1.0.0',
        'timestamp': datetime.datetime.now().isoformat(),
        'whisper_service': app.config['WHISPER_SERVICE_URL']
    })


@app.route('/transcribe_youtube', methods=['POST'])
def transcribe_youtube():
    """Транскрибирование аудио из YouTube видео"""
    try:
        data = request.get_json()
        if not data or 'url' not in data:
            return jsonify({'error': 'URL не указан'}), 400
        
        url = data['url']
        language = data.get('language', 'ru-RU')
        timestamps = data.get('timestamps', False)
        
        # Генерируем ID задачи
        task_id = generate_task_id()
        task_status[task_id] = {
            'status': 'processing',
            'progress': 0,
            'message': 'Подготовка к загрузке видео...'
        }
        
        def update_status(percent, message):
            if task_id in task_status:
                task_status[task_id]['progress'] = percent
                task_status[task_id]['message'] = message
        
        # Запускаем загрузку и транскрипцию в отдельном потоке
        def process_video():
            try:
                # Загружаем видео
                audio_file, video_info = download_from_youtube(url, update_status)
                
                if not audio_file or not os.path.exists(audio_file):
                    raise Exception("Не удалось загрузить аудио из видео")
                
                # Подготавливаем аудио для транскрипции
                prepared_file = prepare_audio_for_transcription(audio_file, update_status)
                
                # Отправляем файл на транскрипцию
                with open(prepared_file, 'rb') as f:
                    files = {'file': (os.path.basename(prepared_file), f, 'audio/wav')}
                    data = {
                        'language': language,
                        'timestamps': str(timestamps).lower()
                    }
                    
                    response = requests.post(
                        f"{app.config['WHISPER_SERVICE_URL']}/transcribe",
                        files=files,
                        data=data
                    )
                    
                    if response.status_code != 200:
                        raise Exception(f"Ошибка API: {response.text}")
                    
                    result = response.json()
                    task_id = result.get('task_id')
                    
                    # Ждем завершения транскрипции
                    while True:
                        status_response = requests.get(
                            f"{app.config['WHISPER_SERVICE_URL']}/status/{task_id}"
                        )
                        status_data = status_response.json()
                        
                        if status_data['status'] == 'completed':
                            task_status[task_id]['status'] = 'completed'
                            task_status[task_id]['progress'] = 100
                            task_status[task_id]['message'] = 'Транскрипция завершена'
                            task_status[task_id]['result'] = status_data['result']
                            break
                        elif status_data['status'] == 'error':
                            raise Exception(f"Ошибка транскрипции: {status_data.get('message', 'Неизвестная ошибка')}")
                        
                        time.sleep(1)
                
            except Exception as e:
                task_status[task_id]['status'] = 'error'
                task_status[task_id]['message'] = str(e)
                print(f"Ошибка при обработке видео: {e}")
                traceback.print_exc()
            finally:
                # Очищаем временные файлы
                try:
                    if 'audio_file' in locals() and os.path.exists(audio_file):
                        os.remove(audio_file)
                    if 'prepared_file' in locals() and os.path.exists(prepared_file):
                        os.remove(prepared_file)
                except Exception as e:
                    print(f"Ошибка при очистке временных файлов: {e}")
        
        # Запускаем обработку в отдельном потоке
        thread = threading.Thread(target=process_video)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'task_id': task_id,
            'message': 'Задача запущена',
            'video_info': {
                'title': 'Загрузка информации...',
                'duration': 0
            }
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)